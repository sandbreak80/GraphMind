"""Process Excel, Word, and text documents with enhanced metadata."""
import logging
from pathlib import Path
from typing import List, Dict, Any
import pandas as pd
from docx import Document
from openpyxl import load_workbook

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Process various document formats."""
    
    def process_excel(self, file_path: Path) -> List[Dict[str, Any]]:
        """
        Extract content from Excel files.
        Extracts data from all sheets, including formulas and comments.
        """
        chunks = []
        doc_id = file_path.stem
        
        try:
            # Read with pandas for data
            excel_file = pd.ExcelFile(file_path)
            
            for sheet_idx, sheet_name in enumerate(excel_file.sheet_names):
                df = excel_file.parse(sheet_name)
                
                # Convert to markdown-like table format
                text_parts = [f"# Sheet: {sheet_name}\n"]
                
                # Add data as formatted text
                if not df.empty:
                    # Get column headers
                    text_parts.append("## Data\n")
                    text_parts.append(" | ".join(str(col) for col in df.columns))
                    text_parts.append("\n" + "-" * 50 + "\n")
                    
                    # Add rows
                    for idx, row in df.iterrows():
                        row_text = " | ".join(str(val) for val in row.values)
                        text_parts.append(row_text + "\n")
                
                # Try to get formulas with openpyxl
                try:
                    wb = load_workbook(file_path, data_only=False)
                    ws = wb[sheet_name]
                    
                    formulas = []
                    comments = []
                    
                    for row in ws.iter_rows():
                        for cell in row:
                            if cell.value and isinstance(cell.value, str) and cell.value.startswith('='):
                                formulas.append(f"{cell.coordinate}: {cell.value}")
                            if cell.comment:
                                comments.append(f"{cell.coordinate}: {cell.comment.text}")
                    
                    if formulas:
                        text_parts.append("\n## Formulas\n")
                        text_parts.extend(f"- {f}\n" for f in formulas)
                    
                    if comments:
                        text_parts.append("\n## Comments\n")
                        text_parts.extend(f"- {c}\n" for c in comments)
                    
                except Exception as e:
                    logger.debug(f"Could not extract formulas from {sheet_name}: {e}")
                
                full_text = "".join(text_parts)
                
                if full_text.strip():
                    chunks.append({
                        "text": full_text,
                        "doc_id": doc_id,
                        "page": sheet_idx + 1,
                        "section": sheet_name,
                        "chunk_id": f"{doc_id}_sheet{sheet_idx}",
                        "content_type": "spreadsheet",
                        "extraction_method": "pandas+openpyxl",
                        "has_formulas": len(formulas) > 0 if 'formulas' in locals() else False,
                        "has_comments": len(comments) > 0 if 'comments' in locals() else False
                    })
            
            logger.info(f"Extracted {len(chunks)} sheets from {file_path.name}")
            return chunks
            
        except Exception as e:
            logger.error(f"Failed to process Excel {file_path}: {e}")
            return []
    
    def process_word(self, file_path: Path) -> List[Dict[str, Any]]:
        """
        Extract content from Word documents.
        Preserves headings, paragraphs, tables, and formatting.
        """
        chunks = []
        doc_id = file_path.stem
        
        try:
            doc = Document(file_path)
            
            text_parts = []
            current_section = "Introduction"
            
            # Extract paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                
                # Detect headings
                if para.style.name.startswith('Heading'):
                    current_section = text
                    text_parts.append(f"\n## {text}\n")
                else:
                    text_parts.append(text + "\n")
            
            # Extract tables
            if doc.tables:
                text_parts.append("\n## Tables\n")
                for table_idx, table in enumerate(doc.tables):
                    text_parts.append(f"\n### Table {table_idx + 1}\n")
                    for row in table.rows:
                        row_text = " | ".join(cell.text.strip() for cell in row.cells)
                        text_parts.append(row_text + "\n")
            
            full_text = "".join(text_parts)
            
            # Chunk the text (simple splitting for now)
            if len(full_text) > 1000:
                # Split into chunks
                words = full_text.split()
                chunk_size = 400
                for i in range(0, len(words), chunk_size):
                    chunk_text = " ".join(words[i:i + chunk_size])
                    chunks.append({
                        "text": chunk_text,
                        "doc_id": doc_id,
                        "page": i // chunk_size + 1,
                        "section": current_section,
                        "chunk_id": f"{doc_id}_c{i // chunk_size}",
                        "content_type": "document",
                        "extraction_method": "python-docx"
                    })
            else:
                chunks.append({
                    "text": full_text,
                    "doc_id": doc_id,
                    "page": 1,
                    "section": current_section,
                    "chunk_id": f"{doc_id}_c0",
                    "content_type": "document",
                    "extraction_method": "python-docx"
                })
            
            logger.info(f"Extracted {len(chunks)} chunks from {file_path.name}")
            return chunks
            
        except Exception as e:
            logger.error(f"Failed to process Word doc {file_path}: {e}")
            return []
    
    def process_text(self, file_path: Path) -> List[Dict[str, Any]]:
        """Extract content from plain text files."""
        try:
            text = file_path.read_text(encoding='utf-8', errors='ignore')
            
            doc_id = file_path.stem
            
            # Simple chunking
            chunks = []
            words = text.split()
            chunk_size = 400
            
            for i in range(0, len(words), chunk_size):
                chunk_text = " ".join(words[i:i + chunk_size])
                if chunk_text.strip():
                    chunks.append({
                        "text": chunk_text,
                        "doc_id": doc_id,
                        "page": i // chunk_size + 1,
                        "section": "Content",
                        "chunk_id": f"{doc_id}_c{i // chunk_size}",
                        "content_type": "text",
                        "extraction_method": "direct"
                    })
            
            logger.info(f"Extracted {len(chunks)} chunks from {file_path.name}")
            return chunks
            
        except Exception as e:
            logger.error(f"Failed to process text file {file_path}: {e}")
            return []
